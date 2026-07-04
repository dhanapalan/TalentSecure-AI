"""
Document Loader & Parser
========================
Handles PDF, DOCX, TXT, and Markdown document ingestion and chunking.
"""

import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

import pypdf as PyPDF2  # pypdf is the maintained successor; same PdfReader API
import pdfplumber
from docx import Document as DocxDocument

from config import config, DOCUMENTS_DIR

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a loaded document"""

    id: str
    source: str
    title: str
    content: str
    metadata: Dict
    chunks: List["DocumentChunk"]


@dataclass
class DocumentChunk:
    """Represents a chunk of document content"""

    id: str
    document_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict


class DocumentLoader:
    """Load and parse documents from various formats"""

    SUPPORTED_FORMATS = ["pdf", "docx", "txt", "md"]

    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        self.chunk_size = chunk_size or config.vector_store.chunk_size
        self.chunk_overlap = chunk_overlap or config.vector_store.chunk_overlap
        self.logger = logger

    def load_document(self, file_path: str) -> Optional[Document]:
        """Load a document from file path"""
        path = Path(file_path)

        if not path.exists():
            self.logger.error(f"File not found: {file_path}")
            return None

        # Check file size
        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > config.document.max_file_size_mb:
            self.logger.error(
                f"File too large: {file_size_mb}MB (max: {config.document.max_file_size_mb}MB)"
            )
            return None

        file_type = path.suffix.lower().lstrip(".")

        if file_type == "pdf":
            return self._load_pdf(path)
        elif file_type == "docx":
            return self._load_docx(path)
        elif file_type in ["txt", "text"]:
            return self._load_txt(path)
        elif file_type == "md":
            return self._load_markdown(path)
        else:
            self.logger.error(f"Unsupported file format: {file_type}")
            return None

    def _load_pdf(self, path: Path) -> Document:
        """Load PDF file"""
        self.logger.info(f"Loading PDF: {path.name}")

        content = ""
        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "pdf",
            "loaded_at": datetime.now().isoformat(),
        }

        try:
            # Try pdfplumber first (better for extraction)
            with pdfplumber.open(path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            content += f"\n\n--- Page {i + 1} ---\n{text}"
                    except Exception as e:
                        self.logger.warning(f"Error extracting page {i + 1}: {e}")

            # Fallback to PyPDF2 if pdfplumber fails
            if not content:
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata["page_count"] = len(reader.pages)

                    for i, page in enumerate(reader.pages):
                        try:
                            text = page.extract_text()
                            if text:
                                content += f"\n\n--- Page {i + 1} ---\n{text}"
                        except Exception as e:
                            self.logger.warning(f"Error extracting page {i + 1}: {e}")

        except Exception as e:
            self.logger.error(f"Error loading PDF: {e}")
            return None

        return Document(
            id=self._generate_id(path),
            source=str(path),
            title=path.stem,
            content=content,
            metadata=metadata,
            chunks=self._chunk_content(content, str(path)),
        )

    def _load_docx(self, path: Path) -> Document:
        """Load DOCX file"""
        self.logger.info(f"Loading DOCX: {path.name}")

        content = ""
        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "docx",
            "loaded_at": datetime.now().isoformat(),
        }

        try:
            doc = DocxDocument(path)
            metadata["paragraph_count"] = len(doc.paragraphs)

            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n"

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content += row_text + "\n"

        except Exception as e:
            self.logger.error(f"Error loading DOCX: {e}")
            return None

        return Document(
            id=self._generate_id(path),
            source=str(path),
            title=path.stem,
            content=content,
            metadata=metadata,
            chunks=self._chunk_content(content, str(path)),
        )

    def _load_txt(self, path: Path) -> Document:
        """Load text file"""
        self.logger.info(f"Loading TXT: {path.name}")

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "txt",
            "loaded_at": datetime.now().isoformat(),
        }

        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception as e:
            self.logger.error(f"Error loading TXT: {e}")
            return None

        return Document(
            id=self._generate_id(path),
            source=str(path),
            title=path.stem,
            content=content,
            metadata=metadata,
            chunks=self._chunk_content(content, str(path)),
        )

    def _load_markdown(self, path: Path) -> Document:
        """Load Markdown file"""
        self.logger.info(f"Loading Markdown: {path.name}")

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "md",
            "loaded_at": datetime.now().isoformat(),
        }

        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception as e:
            self.logger.error(f"Error loading Markdown: {e}")
            return None

        return Document(
            id=self._generate_id(path),
            source=str(path),
            title=path.stem,
            content=content,
            metadata=metadata,
            chunks=self._chunk_content(content, str(path)),
        )

    def _chunk_content(
        self, content: str, source: str, doc_id: str = None
    ) -> List[DocumentChunk]:
        """Split content into overlapping chunks"""
        chunks = []
        content = content.strip()

        for i in range(0, len(content), self.chunk_size - self.chunk_overlap):
            chunk_text = content[i : i + self.chunk_size]

            if len(chunk_text.strip()) > 50:  # Skip very small chunks
                chunk = DocumentChunk(
                    id=f"{self._generate_id(Path(source))}_chunk_{len(chunks)}",
                    document_id=doc_id or self._generate_id(Path(source)),
                    content=chunk_text,
                    chunk_index=len(chunks),
                    start_char=i,
                    end_char=min(i + self.chunk_size, len(content)),
                    metadata={
                        "source": source,
                        "chunk_index": len(chunks),
                        "chunk_size": len(chunk_text),
                    },
                )
                chunks.append(chunk)

        self.logger.info(f"Created {len(chunks)} chunks from document")
        return chunks

    def load_batch(self, directory: str = None) -> List[Document]:
        """Load all supported documents from a directory"""
        directory = directory or config.document.documents_dir
        path = Path(directory)

        if not path.exists():
            self.logger.error(f"Directory not found: {directory}")
            return []

        documents = []

        for file_path in path.glob("**/*"):
            if file_path.is_file() and file_path.suffix.lower().lstrip(".") in self.SUPPORTED_FORMATS:
                doc = self.load_document(str(file_path))
                if doc:
                    documents.append(doc)

        self.logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents

    @staticmethod
    def _generate_id(path: Path) -> str:
        """Generate unique document ID"""
        return f"doc_{path.stem}_{int(path.stat().st_mtime)}".replace(" ", "_").lower()


if __name__ == "__main__":
    # Test document loader
    loader = DocumentLoader()

    # Load a test PDF if it exists
    test_file = DOCUMENTS_DIR / "sample.pdf"
    if test_file.exists():
        doc = loader.load_document(str(test_file))
        if doc:
            print(f"✅ Loaded: {doc.title}")
            print(f"   Chunks: {len(doc.chunks)}")
    else:
        print("No test documents found in", DOCUMENTS_DIR)
